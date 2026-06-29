import asyncio
import base64
import os
import json
import tempfile
import traceback
from collections.abc import AsyncIterator
from pathlib import Path
from typing import Annotated
from dotenv import load_dotenv

import docx as python_docx

load_dotenv()

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
from langgraph_sdk import get_client

from qa_agent.nodes.pdf_extractor import extract_steel_list_pdf, extract_overview_plan_pdf
from qa_agent.pdf_annotator import annotate_pdf
from qa_agent import checks_registry as registry

LANGGRAPH_URL = os.getenv("LANGGRAPH_URL", "http://127.0.0.1:2024")

_KNOWLEDGE_DIR = Path(__file__).parent.parent.parent / "QA AI Drawing" / "QA Knowledge"
_DOCX_PATH     = _KNOWLEDGE_DIR / "structural_qa_rag_knowledge_pack.docx"
_RAG_CACHE_PATH = Path(__file__).parent / "rag" / "data" / "knowledge_cache.json"

# Domain-section keywords (same as knowledge_builder.py)
_DOMAIN_KEYWORDS = {
    "spell": ["spell check node"],
    "bend":  ["bend check node"],
    "rebar": ["rebar check node"],
}
GRAPH_NAME = "qa_agent"

app = FastAPI(title="Drawing Analyzer API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

NODE_LABELS: dict[str, str] = {
    "preprocess":        "Extracting PDF content",
    "spell_check":       "Checking spelling & labels",
    "bend_check":        "Validating bending details",
    "rebar_check":       "Validating rebar specifications",
    "aggregate_results": "Aggregating results",
    "return_to_ui":      "Preparing report",
}

_ALL_CHECKS = ["spell", "bend", "rebar"]


def _sse(payload: dict) -> str:
    return f"data: {json.dumps(payload)}\n\n"


async def _save_upload(data: bytes) -> str:
    """Write upload bytes to a temp file without blocking the event loop."""
    def _write() -> str:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(data)
            return tmp.name
    return await asyncio.to_thread(_write)


def _node_sse_events(node_name: str, node_data: dict) -> list[str]:
    """Convert a completed node into SSE event strings."""
    label = NODE_LABELS.get(node_name, node_name)
    events = [_sse({"type": "progress", "node": node_name, "label": label})]
    print(f"[server] ✓ node: {node_name}  |  keys: {list(node_data.keys()) if isinstance(node_data, dict) else type(node_data)}")

    if node_name != "return_to_ui":
        return events

    ui_response = node_data.get("ui_response") if isinstance(node_data, dict) else None
    print(f"[server] ui_response summary: {ui_response.get('summary') if ui_response else 'NONE'}")
    if ui_response:
        events.append(_sse({"type": "result", "data": ui_response}))
    return events


async def _run_graph(
    tmp_path: str,
    enabled_checks: list[str],
    enabled_sub_checks: dict[str, list[str]] | None = None,
    steel_list_data: dict | None = None,
    overview_plan_data: dict | None = None,
) -> AsyncIterator[str]:
    """Stream SSE events by routing graph execution through LangGraph API."""
    lg = get_client(url=LANGGRAPH_URL)
    thread = await lg.threads.create()
    thread_id = thread["thread_id"]
    studio_url = f"https://smith.langchain.com/studio/?baseUrl={LANGGRAPH_URL}"
    print(f"[server] thread_id : {thread_id}")
    print(f"[server] studio    : {studio_url}")

    yield _sse({"type": "run_started", "thread_id": thread_id, "studio_url": studio_url})

    graph_input: dict = {"pdf_path": tmp_path, "enabled_checks": enabled_checks}
    if enabled_sub_checks:
        graph_input["enabled_sub_checks"] = enabled_sub_checks
    if steel_list_data:
        graph_input["steel_list_data"] = steel_list_data
    if overview_plan_data:
        graph_input["overview_plan_data"] = overview_plan_data

    async for chunk in lg.runs.stream(
        thread_id,
        GRAPH_NAME,
        input=graph_input,
        stream_mode="updates",
    ):
        if chunk.event == "metadata":
            run_id = chunk.data.get("run_id") if isinstance(chunk.data, dict) else None
            if run_id:
                print(f"[server] run_id    : {run_id}")
        elif chunk.event == "updates":
            for node_name, node_data in chunk.data.items():
                for event in _node_sse_events(node_name, node_data):
                    yield event
        elif chunk.event == "error":
            print(f"[server] ✗ LangGraph error: {chunk.data}")
            yield _sse({"type": "error", "message": str(chunk.data)})


@app.post("/api/analyze")
async def analyze(
    file: Annotated[UploadFile, File()],
    checks: Annotated[str, Form()] = "spell,bend,rebar",
    sub_checks: Annotated[str, Form()] = "{}",
    steel_list: Annotated[UploadFile | None, File()] = None,
    overview_plan: Annotated[UploadFile | None, File()] = None,
):
    enabled_checks = [c.strip() for c in checks.split(",") if c.strip() in _ALL_CHECKS] or _ALL_CHECKS[:]

    try:
        enabled_sub_checks: dict[str, list[str]] = json.loads(sub_checks) if sub_checks else {}
    except (json.JSONDecodeError, ValueError):
        enabled_sub_checks = {}

    data = await file.read()
    tmp_path = await _save_upload(data)

    print(f"\n[server] === New analysis: {file.filename} → {tmp_path} | checks: {enabled_checks} | sub_checks: {enabled_sub_checks} ===")
    print(f"[server] Routing to LangGraph API at {LANGGRAPH_URL}")

    # ── pdfplumber pre-extraction for supplementary files ───────────────────
    steel_list_data: dict | None = None
    overview_plan_data: dict | None = None

    if steel_list and steel_list.filename:
        sl_bytes = await steel_list.read()
        sl_tmp = await _save_upload(sl_bytes)
        try:
            steel_list_data = await asyncio.to_thread(extract_steel_list_pdf, sl_tmp)
            # Carry the raw PDF (base64) so checks can read the table from the
            # rendered document (vision) — pdfplumber text loses cells like the
            # Korrosionsschutz "FV" code that sit in narrow/wrapped columns.
            steel_list_data["pdf_data"] = base64.b64encode(sl_bytes).decode("ascii")
            print(f"[server] steel_list extracted: gesamtmasse={steel_list_data.get('gesamtmasse')!r}  pages={steel_list_data.get('page_count')}")
        except Exception as exc:
            print(f"[server] ✗ steel_list extraction failed: {exc}")
        finally:
            try:
                os.unlink(sl_tmp)
            except OSError:
                pass

    if overview_plan and overview_plan.filename:
        op_bytes = await overview_plan.read()
        op_tmp = await _save_upload(op_bytes)
        try:
            overview_plan_data = await asyncio.to_thread(extract_overview_plan_pdf, op_tmp)
            print(f"[server] overview_plan extracted: pages={overview_plan_data.get('page_count')}  chars={len(overview_plan_data.get('raw_text', ''))}")
        except Exception as exc:
            print(f"[server] ✗ overview_plan extraction failed: {exc}")
        finally:
            try:
                os.unlink(op_tmp)
            except OSError:
                pass

    async def stream():
        yield _sse({"type": "ack"})
        try:
            async for event in _run_graph(tmp_path, enabled_checks, enabled_sub_checks, steel_list_data, overview_plan_data):
                yield event
        except Exception as exc:
            print(f"[server] ✗ EXCEPTION: {exc}")
            traceback.print_exc()
            yield _sse({"type": "error", "message": str(exc)})
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
            print("[server] === Analysis done ===\n")

    return StreamingResponse(stream(), media_type="text/event-stream")


@app.post("/api/annotate")
async def annotate_report(
    file: Annotated[UploadFile, File()],
    result: Annotated[str, Form()],
):
    """Return the uploaded PDF annotated with the failed findings from `result`.

    `result` is the analysis-result JSON the frontend already holds. Only real
    findings (ERROR/WARNING, not per-check summaries or not-found items) are
    annotated — each anchored where its offending text appears in the drawing.
    """
    pdf_bytes = await file.read()
    try:
        parsed = json.loads(result) if result else {}
    except (json.JSONDecodeError, ValueError):
        raise HTTPException(status_code=400, detail="Invalid result JSON")

    sections = parsed.get("sections") if isinstance(parsed, dict) else None
    issues: list[dict] = []
    for section in (sections or []):
        if not isinstance(section, dict):
            continue
        current_check = section.get("title")
        for it in (section.get("issues") or []):
            if not isinstance(it, dict):
                continue
            # Per-check summary / not-found item — remember its name, don't annotate.
            if "passed" in it or it.get("not_found") is True:
                current_check = it.get("check_name") or current_check
                continue
            sev = str(it.get("severity", "")).upper()
            if sev not in ("ERROR", "WARNING"):
                continue
            issues.append({
                "page":        it.get("page", 1),
                "description": it.get("description", ""),
                "location":    it.get("location", ""),
                "severity":    sev,
                "check_name":  it.get("check_name") or current_check or it.get("category"),
                "category":    it.get("category"),
            })

    print(f"[annotate] {len(issues)} failed finding(s) → annotating {file.filename!r}")
    try:
        annotated = await asyncio.to_thread(annotate_pdf, pdf_bytes, issues)
    except Exception as exc:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Annotation failed: {exc}")

    base = (file.filename or "drawing.pdf").rsplit(".", 1)[0]
    return Response(
        content=annotated,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{base}_annotated.pdf"'},
    )


@app.get("/api/extraction-fields")
async def extraction_fields():
    """Catalog of pdfplumber fields available to user-defined checks (for frontend)."""
    from qa_agent.extraction import list_extraction_fields
    return {"fields": await asyncio.to_thread(list_extraction_fields)}


@app.get("/api/checks")
async def list_checks():
    """All defined checks (built-in + custom), for the Define-Rules UI and the
    dynamic check-selection sidebar."""
    checks = await asyncio.to_thread(registry.list_checks)
    domains = [
        {
            "key": d,
            "title": registry.domain_title(d),
            "coming_soon": d in registry.COMING_SOON_DOMAINS,
        }
        for d in registry.ALL_DOMAINS
    ]
    return {"checks": checks, "domains": domains}


class _CheckBody(BaseModel):
    domain: str = "spell"
    key: str | None = None
    display_name: str
    description: str = ""
    prompt: str = ""
    pass_text: str = "PASS"
    not_found_text: str = "NOT FOUND"
    requires_vision: bool = False


@app.post("/api/checks")
async def save_check(body: _CheckBody):
    try:
        saved = await asyncio.to_thread(
            registry.save_check,
            domain=body.domain, key=body.key, display_name=body.display_name,
            description=body.description, prompt=body.prompt,
            pass_text=body.pass_text, not_found_text=body.not_found_text,
            requires_vision=body.requires_vision,
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))
    return {"ok": True, "check": saved}


@app.delete("/api/checks/{domain}/{key}")
async def delete_check(domain: str, key: str):
    try:
        await asyncio.to_thread(registry.delete_check, domain, key)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))
    return {"ok": True}


@app.get("/api/knowledge-base/download")
async def download_knowledge_file():
    def _ensure() -> None:
        _KNOWLEDGE_DIR.mkdir(parents=True, exist_ok=True)
        if not _DOCX_PATH.exists():
            doc = python_docx.Document()
            doc.add_heading("QA Knowledge Base", level=1)
            for heading in ("Spell Check Node", "Bend Check Node", "Rebar Check Node"):
                doc.add_heading(heading, level=2)
                doc.add_paragraph("Add your QA rules here.")
            doc.save(str(_DOCX_PATH))

    await asyncio.to_thread(_ensure)
    return FileResponse(
        path=str(_DOCX_PATH),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="structural_qa_rag_knowledge_pack.docx",
    )


_SUPPORTED_IMAGE_TYPES = {"image/png", "image/jpeg", "image/gif", "image/webp"}
_R_EMBED = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"


@app.post("/api/knowledge-base/upload")
async def upload_knowledge_file(file: Annotated[UploadFile, File()]):
    import base64  # noqa: PLC0415

    data = await file.read()

    def _save_and_index() -> tuple[int, int]:
        _KNOWLEDGE_DIR.mkdir(parents=True, exist_ok=True)
        _DOCX_PATH.write_bytes(data)

        doc = python_docx.Document(str(_DOCX_PATH))

        # ── Build rId → image map from document relationships ──────────────
        rid_to_img: dict[str, dict] = {}
        for rel in doc.part.rels.values():
            if "image" not in rel.reltype:
                continue
            media_type = rel.target_part.content_type
            if media_type not in _SUPPORTED_IMAGE_TYPES:
                continue
            b64 = base64.standard_b64encode(rel.target_part.blob).decode()
            rid_to_img[rel.rId] = {"data": b64, "media_type": media_type}

        # ── Walk paragraphs: extract text + images, split by domain ────────
        text_sections: dict[str, list[str]] = {d: [] for d in ("spell", "bend", "rebar")}
        img_sections:  dict[str, list[dict]] = {d: [] for d in ("spell", "bend", "rebar")}
        current: str | None = None

        for para in doc.paragraphs:
            text = para.text.strip()
            lower = text.lower()
            for domain, keywords in _DOMAIN_KEYWORDS.items():
                if any(kw in lower for kw in keywords):
                    current = domain
                    break
            if not current:
                continue
            if text:
                text_sections[current].append(text)
            for elem in para._element.iter():
                rid = elem.get(_R_EMBED)
                if rid and rid in rid_to_img:
                    img = rid_to_img[rid]
                    if img not in img_sections[current]:   # deduplicate
                        img_sections[current].append(img)

        # Fall back: no domain headers found → assign everything to all domains
        if not any(lines for lines in text_sections.values()):
            all_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            all_imgs  = list(rid_to_img.values())
            for domain in ("spell", "bend", "rebar"):
                text_sections[domain] = all_text
                img_sections[domain]  = all_imgs

        cache: dict = {}
        if _RAG_CACHE_PATH.exists():
            with open(_RAG_CACHE_PATH, encoding="utf-8") as f:
                cache = json.load(f)
        cache["docx_knowledge"] = {d: "\n".join(lines) for d, lines in text_sections.items()}
        cache["docx_images"]    = img_sections
        _RAG_CACHE_PATH.parent.mkdir(parents=True, exist_ok=True)
        with open(_RAG_CACHE_PATH, "w", encoding="utf-8") as f:
            json.dump(cache, f, indent=2, ensure_ascii=False)

        total_imgs = sum(len(v) for v in img_sections.values())
        return sum(len(v) for v in text_sections.values()), total_imgs

    lines, imgs = await asyncio.to_thread(_save_and_index)
    from qa_agent.rag.retriever import _load_cache  # noqa: PLC0415
    _load_cache.cache_clear()
    return {"ok": True, "lines_indexed": lines, "images_indexed": imgs}
